"""Resume export endpoints — PDF, DOCX, JSON generation."""

from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel
from typing import Optional
import io

from sqlalchemy.ext.asyncio import AsyncSession

from app.core.database import get_db
from app.crud.resume import get_resume
from app.utils.dependencies import get_current_user

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

router = APIRouter(prefix="/exports", tags=["resume-exports"])


class PDFExportRequest(BaseModel):
    resume_id: str
    template_id: str = "meridian"
    paper_size: str = "letter"  # "letter" | "a4"


class PDFExportResponse(BaseModel):
    job_id: str
    status: str = "queued"
    message: str = "PDF generation has been queued"


class ExportStatusResponse(BaseModel):
    job_id: str
    status: str  # "queued" | "processing" | "completed" | "failed"
    file_url: Optional[str] = None
    error: Optional[str] = None


class DOCXExportRequest(BaseModel):
    resume_data: dict
    template_id: str = "meridian"


class TemplateInfo(BaseModel):
    id: str
    name: str
    description: str
    best_for: str
    preview_url: Optional[str] = None


# ── PDF Export ───────────────────────────────────────────────


@router.post("/pdf", response_model=PDFExportResponse)
async def export_pdf(
    request: PDFExportRequest,
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Queue a PDF export job for one of the caller's OWN resumes.

    Auth is REQUIRED (it was commented out "for dev"). We also verify the
    resume belongs to the caller before queueing, so an authenticated user
    can't enqueue exports of someone else's resume by guessing IDs.
    """
    # Ownership gate: get_resume returns None unless the resume exists AND
    # belongs to current_user.
    resume = await get_resume(db, request.resume_id, current_user.id)
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    # NOTE: the BullMQ + Puppeteer render worker is a future integration not
    # yet part of this Python service — the job is accepted but no worker
    # completes it. /exports/docx is a fully-working real download today.
    import uuid
    job_id = str(uuid.uuid4())

    return PDFExportResponse(
        job_id=job_id,
        status="queued",
        message="PDF generation queued. Poll /exports/status/{job_id} for completion.",
    )


@router.get("/status/{job_id}", response_model=ExportStatusResponse)
async def export_status(
    job_id: str,
    current_user=Depends(get_current_user),
):
    """Poll the status of a PDF export job. Auth REQUIRED (was disabled).

    STUB: there's no BullMQ jobs table to look the id up in yet, so this
    returns a placeholder. Wire the worker + a jobs table to make it real
    (and to scope job lookups to current_user.id).
    """
    return ExportStatusResponse(
        job_id=job_id,
        status="completed",
        file_url=f"/api/v1/exports/download/{job_id}",
    )


# ── DOCX Export ──────────────────────────────────────────────


@router.post("/docx")
async def export_docx(
    request: DOCXExportRequest,
    current_user=Depends(get_current_user),
):
    """Generate and return a DOCX file from resume data (auth REQUIRED).

    The caller supplies the resume_data in the body, so there's no ID to
    ownership-check — but it must not be an open, unauthenticated render
    service.
    """
    if not HAS_DOCX:
        raise HTTPException(status_code=503, detail="python-docx not installed. Run: pip install python-docx")

    data = request.resume_data
    doc = DocxDocument()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    header = data.get("header", {})

    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(f"{header.get('firstName', '')} {header.get('lastName', '')}")
    name_run.bold = True
    name_run.font.size = Pt(18)

    # Headline
    if header.get("headline"):
        headline_para = doc.add_paragraph()
        headline_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = headline_para.add_run(header["headline"])
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(100, 100, 100)

    # Contact line
    contact = [header.get("email", ""), header.get("phone", ""), header.get("location", "")]
    contact = [c for c in contact if c]
    if contact:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run(" | ".join(contact))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(120, 120, 120)

    links = [header.get("linkedin", ""), header.get("github", ""), header.get("website", "")]
    links = [l for l in links if l]
    if links:
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = lp.add_run(" | ".join(links))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(120, 120, 120)

    # Summary
    summary = data.get("summary", {}).get("content", "")
    if summary:
        doc.add_heading("PROFESSIONAL SUMMARY", level=2)
        p = doc.add_paragraph(summary)
        p.style.font.size = Pt(10)

    # Experience
    experience = data.get("experience", [])
    if experience:
        doc.add_heading("EXPERIENCE", level=2)
        for exp in experience:
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(exp.get("title", ""))
            title_run.bold = True
            title_run.font.size = Pt(11)
            title_para.add_run(f" — {exp.get('company', '')}")

            date_para = doc.add_paragraph()
            date_run = date_para.add_run(
                f"{exp.get('startDate', '')} — {'Present' if exp.get('current') else exp.get('endDate', '')}"
                + (f" | {exp.get('location', '')}" if exp.get('location') else "")
            )
            date_run.font.size = Pt(9)
            date_run.font.color.rgb = RGBColor(120, 120, 120)

            for bullet in exp.get("bullets", []):
                content = bullet.get("content", "") if isinstance(bullet, dict) else bullet
                if content:
                    bp = doc.add_paragraph(content, style="List Bullet")
                    if isinstance(bullet, dict) and bullet.get("highlighted"):
                        for run in bp.runs:
                            run.bold = True

    # Education
    education = data.get("education", [])
    if education:
        doc.add_heading("EDUCATION", level=2)
        for edu in education:
            p = doc.add_paragraph()
            run = p.add_run(f"{edu.get('degree', '')}{' in ' + edu.get('field', '') if edu.get('field') else ''}")
            run.bold = True
            run.font.size = Pt(11)
            p2 = doc.add_paragraph()
            run2 = p2.add_run(f"{edu.get('institution', '')} — {edu.get('endDate', '')}")
            run2.font.size = Pt(10)
            run2.italic = True
            if edu.get("gpa"):
                doc.add_paragraph(f"GPA: {edu['gpa']}")

    # Skills
    skills = data.get("skills", {}).get("groups", [])
    if any(g.get("skills") for g in skills):
        doc.add_heading("SKILLS", level=2)
        for group in skills:
            if group.get("skills"):
                p = doc.add_paragraph()
                run = p.add_run(f"{group.get('label', 'Skills')}: ")
                run.bold = True
                p.add_run(", ".join(group["skills"]))

    # Projects
    projects = data.get("projects", [])
    if projects:
        doc.add_heading("PROJECTS", level=2)
        for proj in projects:
            p = doc.add_paragraph()
            run = p.add_run(proj.get("name", ""))
            run.bold = True
            if proj.get("description"):
                p.add_run(f" — {proj['description']}")
            for bullet in proj.get("bullets", []):
                content = bullet.get("content", "") if isinstance(bullet, dict) else bullet
                if content:
                    doc.add_paragraph(content, style="List Bullet")

    # Certifications
    certs = data.get("certifications", [])
    if certs:
        doc.add_heading("CERTIFICATIONS", level=2)
        for cert in certs:
            p = doc.add_paragraph()
            run = p.add_run(cert.get("name", ""))
            run.bold = True
            if cert.get("issuer"):
                p.add_run(f" — {cert['issuer']}")
            if cert.get("date"):
                p.add_run(f" ({cert['date']})")

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    first = header.get("firstName", "Resume")
    last = header.get("lastName", "")
    filename = f"{first}-{last}-Resume.docx" if last else f"{first}-Resume.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── JSON Export ──────────────────────────────────────────────


@router.get("/json/{resume_id}")
async def export_json(
    resume_id: str,
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Download one of the caller's OWN resumes as a JSON file.

    Auth REQUIRED + ownership-enforced (was commented out, and the body was
    a placeholder — had the DB fetch been wired in without auth, this would
    have let anyone download any resume by ID). get_resume only returns the
    row when it belongs to current_user, so a non-owner gets a 404.
    """
    resume = await get_resume(db, resume_id, current_user.id)
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    payload = {
        "id": str(resume.id),
        "title": resume.title,
        "template_id": resume.template_id,
        "data": resume.data,
    }
    filename = f"{(resume.title or 'resume').replace(' ', '-')}.json"
    return JSONResponse(
        content=payload,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Templates ────────────────────────────────────────────────


TEMPLATES = [
    TemplateInfo(id="meridian", name="Meridian", description="Clean, single-column with bold headers", best_for="Tech / Engineering"),
    TemplateInfo(id="atlas", name="Atlas", description="Two-column with sidebar for skills and contact", best_for="Business / Finance"),
    TemplateInfo(id="prism", name="Prism", description="Colorful header band with modern typography", best_for="Design / Marketing"),
    TemplateInfo(id="scholar", name="Scholar", description="Academic, dense, publications-ready", best_for="Academia / Research"),
    TemplateInfo(id="carta", name="Carta", description="Minimalist with generous whitespace", best_for="Executives / Senior"),
    TemplateInfo(id="vertex", name="Vertex", description="Technical grid layout with clear hierarchy", best_for="Engineers / PMs"),
    TemplateInfo(id="foundry", name="Foundry", description="Dark header, bold typography", best_for="Creative / Startups"),
    TemplateInfo(id="pulse", name="Pulse", description="Timeline-based with visual progression", best_for="UX / Product"),
]


@router.get("/templates", response_model=list[TemplateInfo])
async def list_templates():
    """List all available resume templates."""
    return TEMPLATES


@router.get("/templates/{template_id}", response_model=TemplateInfo)
async def get_template(template_id: str):
    """Get metadata for a specific template."""
    template = next((t for t in TEMPLATES if t.id == template_id), None)
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    return template
